"""
Serviço de OCR (Reconhecimento Óptico de Caracteres).

Responsável por converter PDFs (escaneados ou digitais) em texto,
preparando o conteúdo para o parser de extratos.

Dependências externas:
  - Tesseract OCR instalado no sistema (https://github.com/UB-Mannheim/tesseract/wiki)
  - pdf2image (converte páginas PDF em imagens)
  - pytesseract (wrapper Python para o Tesseract)
  - python-docx (leitura de arquivos DOCX/DOC)
"""

import io
import os
import re
import logging
from pathlib import Path
from typing import Optional

try:
    import pytesseract
    from PIL import Image, ImageFilter, ImageEnhance
    PYTESSERACT_OK = True
except ImportError:
    PYTESSERACT_OK = False

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_OK = True
except ImportError:
    PDF2IMAGE_OK = False

try:
    import docx as python_docx
    DOCX_OK = True
except ImportError:
    DOCX_OK = False


class DocumentoProtegidoError(Exception):
    """
    Lançada quando o arquivo está protegido por senha.
    Carrega o 'tipo_arquivo' ("pdf" ou "docx") para que o chamador
    possa solicitar a senha ao usuário e tentar novamente.
    """
    def __init__(self, caminho: str, tipo_arquivo: str = "pdf"):
        super().__init__(f"Documento protegido por senha: {caminho}")
        self.caminho      = caminho
        self.tipo_arquivo = tipo_arquivo

logger = logging.getLogger(__name__)


class OCRService:
    """
    Serviço de OCR para extração de texto de PDFs e imagens.

    Suporta:
      - PDFs digitais (extração direta de texto)
      - PDFs escaneados (conversão para imagem + Tesseract)
      - Imagens (PNG, JPG, JPEG, BMP, TIFF, WEBP)
      - Documentos Word (DOCX)
      - Texto plano (TXT)
    """

    # Formatos de imagem suportados
    FORMATOS_IMAGEM = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif", ".webp", ".gif"}

    # Caminho padrão do Tesseract no Windows
    TESSERACT_PATHS_WINDOWS = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]

    def __init__(self, tesseract_path: Optional[str] = None, lang: str = "por"):
        """
        Inicializa o serviço de OCR.

        Args:
            tesseract_path: Caminho manual para o executável do Tesseract
            lang:           Idioma principal para OCR (padrão: "por" = português)
        """
        self._configurar_tesseract(tesseract_path)
        self.lang = self._resolver_idioma_ocr(lang)

    # --------------------------------------------------
    # Configuração
    # --------------------------------------------------

    def _configurar_tesseract(self, caminho_manual: Optional[str]) -> None:
        """Localiza e configura o executável do Tesseract."""
        if not PYTESSERACT_OK:
            logger.warning("pytesseract não instalado. OCR desativado.")
            return

        if caminho_manual and os.path.exists(caminho_manual):
            pytesseract.pytesseract.tesseract_cmd = caminho_manual
            return

        # Tenta caminhos padrão no Windows
        for caminho in self.TESSERACT_PATHS_WINDOWS:
            if os.path.exists(caminho):
                pytesseract.pytesseract.tesseract_cmd = caminho
                logger.info(f"Tesseract encontrado em: {caminho}")
                return

        logger.warning(
            "Tesseract não encontrado nos caminhos padrão. "
            "Instale em https://github.com/UB-Mannheim/tesseract/wiki"
        )

    def _resolver_idioma_ocr(self, idioma_preferido: str) -> str:
        """Seleciona um idioma OCR disponível, com fallback seguro."""
        if not PYTESSERACT_OK:
            return idioma_preferido

        try:
            idiomas_disponiveis = set(pytesseract.get_languages(config=""))
        except Exception as exc:
            logger.warning("Não foi possível listar idiomas do Tesseract: %s", exc)
            return idioma_preferido

        if idioma_preferido in idiomas_disponiveis:
            return idioma_preferido

        if idioma_preferido == "por" and "eng" in idiomas_disponiveis:
            logger.warning(
                "Idioma OCR 'por' não encontrado. Usando fallback 'eng'."
            )
            return "eng"

        if idiomas_disponiveis:
            idioma_fallback = "eng" if "eng" in idiomas_disponiveis else sorted(idiomas_disponiveis)[0]
            logger.warning(
                "Idioma OCR '%s' não encontrado. Usando '%s'.",
                idioma_preferido,
                idioma_fallback,
            )
            return idioma_fallback

        return idioma_preferido

    # --------------------------------------------------
    # Interface pública
    # --------------------------------------------------

    def extrair_texto_pdf(self, caminho_pdf: str, senha: Optional[str] = None) -> str:
        """
        Extrai texto de um arquivo PDF.

        Tenta extração digital primeiro (mais precisa).
        Se o resultado for insuficiente, usa OCR com Tesseract.
        Lança DocumentoProtegidoError se o PDF estiver protegido por senha
        e nenhuma senha (ou senha incorreta) for fornecida.

        Args:
            caminho_pdf: Caminho absoluto do arquivo PDF
            senha:       Senha de abertura (None = sem senha)

        Returns:
            Texto extraído do PDF
        """
        if not os.path.exists(caminho_pdf):
            raise FileNotFoundError(f"Arquivo não encontrado: {caminho_pdf}")

        # Detecta proteção por senha antes de tentar extrair
        self._verificar_protecao_pdf(caminho_pdf, senha)

        # Tenta extração digital via PyMuPDF/pdfplumber (sem OCR)
        texto = self._extrair_texto_digital(caminho_pdf, senha=senha)

        # Se o texto for muito curto, provavelmente é um PDF escaneado
        if len(texto.strip()) < 50:
            logger.info("PDF parece escaneado. Usando OCR com Tesseract...")
            texto = self._extrair_texto_ocr(caminho_pdf, senha=senha)

        return texto

    def extrair_texto_imagem(self, caminho_imagem: str) -> str:
        """
        Extrai texto de um arquivo de imagem (PNG, JPG, JPEG, BMP, TIFF, WEBP).

        Args:
            caminho_imagem: Caminho da imagem

        Returns:
            Texto extraído pela OCR
        """
        if not PYTESSERACT_OK:
            raise RuntimeError("pytesseract não está instalado.")

        imagem = Image.open(caminho_imagem)
        imagem = self._preprocessar_imagem(imagem)
        return pytesseract.image_to_string(imagem, lang=self.lang)

    def extrair_texto_docx(self, caminho_docx: str, senha: Optional[str] = None) -> str:
        """
        Extrai texto de um arquivo Word (.docx).
        Lança DocumentoProtegidoError se o arquivo estiver criptografado
        e nenhuma senha for fornecida.

        Args:
            caminho_docx: Caminho do arquivo DOCX
            senha:        Senha de abertura (None = sem senha)

        Returns:
            Texto extraído do documento
        """
        if not DOCX_OK:
            raise RuntimeError(
                "python-docx não instalado. "
                "Instale com: pip install python-docx"
            )
        if not os.path.exists(caminho_docx):
            raise FileNotFoundError(f"Arquivo não encontrado: {caminho_docx}")

        try:
            doc = python_docx.Document(caminho_docx)
        except Exception as e:
            msg = str(e).lower()
            # python-docx lança exceção ao abrir DOCX criptografado
            if any(kw in msg for kw in ("encrypted", "password", "criptograf", "protegido", "badzip")):
                raise DocumentoProtegidoError(caminho_docx, "docx") from e
            raise

        paragrafos = [p.text for p in doc.paragraphs if p.text.strip()]
        # Inclui texto de tabelas
        for tabela in doc.tables:
            for linha in tabela.rows:
                celulas = [c.text.strip() for c in linha.cells if c.text.strip()]
                if celulas:
                    paragrafos.append("  ".join(celulas))
        return "\n".join(paragrafos)

    def extrair_texto_txt(self, caminho_txt: str) -> str:
        """
        Lê o conteúdo de um arquivo de texto plano (.txt, .csv, etc.).

        Args:
            caminho_txt: Caminho do arquivo

        Returns:
            Conteúdo do arquivo como string
        """
        if not os.path.exists(caminho_txt):
            raise FileNotFoundError(f"Arquivo não encontrado: {caminho_txt}")

        for enc in ("utf-8", "latin-1", "utf-8-sig", "cp1252"):
            try:
                with open(caminho_txt, "r", encoding=enc) as f:
                    return f.read()
            except (UnicodeDecodeError, LookupError):
                continue
        # último recurso: ignora erros de encoding
        with open(caminho_txt, "r", encoding="utf-8", errors="replace") as f:
            return f.read()

    def extrair_texto(self, caminho: str, senha: Optional[str] = None) -> str:
        """
        Dispatcher universal: extrai texto de qualquer formato suportado
        com base na extensão do arquivo.

        Formatos suportados: PDF, JPG, JPEG, PNG, BMP, TIFF, WEBP, DOCX, TXT
        Lança DocumentoProtegidoError se o arquivo necessitar de senha.

        Args:
            caminho: Caminho do arquivo
            senha:   Senha de abertura (None = sem senha)

        Returns:
            Texto extraído
        """
        ext = Path(caminho).suffix.lower()

        if ext == ".pdf":
            return self.extrair_texto_pdf(caminho, senha=senha)
        if ext in self.FORMATOS_IMAGEM:
            return self.extrair_texto_imagem(caminho)
        if ext in (".docx", ".doc"):
            return self.extrair_texto_docx(caminho, senha=senha)
        if ext in (".txt", ".text"):
            return self.extrair_texto_txt(caminho)

        # Fallback: tenta ler como texto plano
        logger.warning(f"Extensão '{ext}' não reconhecida. Tentando leitura como texto.")
        return self.extrair_texto_txt(caminho)

    # --------------------------------------------------
    # Métodos internos
    # --------------------------------------------------

    # --------------------------------------------------
    # Detecção e abertura de PDFs protegidos
    # --------------------------------------------------

    def _verificar_protecao_pdf(self, caminho_pdf: str, senha: Optional[str]) -> None:
        """
        Verifica se o PDF está protegido por senha.
        Lança DocumentoProtegidoError se estiver criptografado e a senha
        fornecida estiver errada ou ausente.
        """
        try:
            import fitz  # PyMuPDF — mais confiável para detectar criptografia
            doc = fitz.open(caminho_pdf)
            if doc.is_encrypted:
                if senha is None:
                    doc.close()
                    raise DocumentoProtegidoError(caminho_pdf, "pdf")
                ok = doc.authenticate(senha)
                doc.close()
                if not ok:
                    raise ValueError(
                        "Senha incorreta para o arquivo PDF. Tente novamente."
                    )
            else:
                doc.close()
            return
        except ImportError:
            pass

        # Fallback: pdfplumber
        try:
            import pdfplumber
            with pdfplumber.open(caminho_pdf, password=senha or "") as pdf:
                _ = pdf.pages  # força abertura
        except Exception as e:
            msg = str(e).lower()
            if any(kw in msg for kw in ("password", "encrypted", "incorrect", "criptograf")):
                if senha is None:
                    raise DocumentoProtegidoError(caminho_pdf, "pdf") from e
                raise ValueError("Senha incorreta para o arquivo PDF. Tente novamente.") from e

    def _extrair_texto_digital(self, caminho_pdf: str, senha: Optional[str] = None) -> str:
        """
        Tenta extrair texto de um PDF digital (sem escaneamento).
        Usa pdfplumber se disponível, senão retorna string vazia.
        """
        try:
            import pdfplumber
            texto_completo = []
            kwargs = {"password": senha} if senha else {}
            with pdfplumber.open(caminho_pdf, **kwargs) as pdf:
                for i, pagina in enumerate(pdf.pages):
                    if i >= 100:
                        logger.warning("PDF com mais de 100 páginas; extração limitada a 100.")
                        break
                    try:
                        t = pagina.extract_text()
                        if t:
                            texto_completo.append(t)
                    except Exception as e_pag:
                        logger.warning("Erro ao extrair página %d: %s", i + 1, e_pag)
            return "\n".join(texto_completo)
        except ImportError:
            pass
        except Exception as e_plumber:
            logger.warning("pdfplumber falhou: %s; tentando PyMuPDF...", e_plumber)

        try:
            import fitz  # PyMuPDF
            doc = fitz.open(caminho_pdf)
            if doc.is_encrypted and senha:
                doc.authenticate(senha)
            texto_completo = []
            for i, pagina in enumerate(doc):
                if i >= 100:
                    logger.warning("PDF com mais de 100 páginas; extração limitada a 100 (fitz).")
                    break
                try:
                    texto_completo.append(pagina.get_text())
                except Exception as e_pag:
                    logger.warning("Erro ao extrair página %d (fitz): %s", i + 1, e_pag)
            doc.close()
            return "\n".join(texto_completo)
        except ImportError:
            pass
        except Exception as e_fitz:
            logger.warning("PyMuPDF falhou: %s", e_fitz)

        return ""

    def _extrair_texto_ocr(self, caminho_pdf: str, senha: Optional[str] = None) -> str:
        """
        Converte cada página do PDF em imagem e aplica OCR.

        Args:
            caminho_pdf: Caminho do arquivo PDF
            senha:       Senha de abertura (se protegido)

        Returns:
            Texto concatenado de todas as páginas
        """
        if not PYTESSERACT_OK:
            raise RuntimeError("pytesseract não instalado.")

        paginas = []
        erro_pdf2image = None

        if PDF2IMAGE_OK:
            kwargs_conv = {"userpw": senha} if senha else {}
            try:
                paginas = convert_from_path(caminho_pdf, dpi=300, **kwargs_conv)
            except Exception as exc:
                erro_pdf2image = exc
                logger.warning(
                    "pdf2image falhou ao rasterizar PDF; tentando PyMuPDF: %s",
                    exc,
                )

        if not paginas:
            paginas = self._renderizar_paginas_pdf_com_fitz(caminho_pdf, senha=senha)

        if not paginas:
            if erro_pdf2image is not None:
                raise RuntimeError(
                    "Não foi possível rasterizar o PDF para OCR. "
                    "Instale o Poppler ou mantenha PyMuPDF disponível."
                ) from erro_pdf2image
            raise RuntimeError(
                "Não foi possível rasterizar o PDF para OCR. "
                "Instale o Poppler ou mantenha PyMuPDF disponível."
            )

        textos = []

        for i, pagina in enumerate(paginas):
            logger.debug(f"Processando página {i + 1}/{len(paginas)}...")
            img_processada = self._preprocessar_imagem(pagina)
            texto = pytesseract.image_to_string(img_processada, lang=self.lang)
            textos.append(texto)

        return "\n".join(textos)

    def _renderizar_paginas_pdf_com_fitz(self, caminho_pdf: str, senha: Optional[str] = None):
        """Renderiza páginas do PDF em imagens PIL usando PyMuPDF."""
        try:
            import fitz
        except ImportError:
            return []

        imagens = []
        doc = fitz.open(caminho_pdf)
        try:
            if doc.is_encrypted and senha:
                doc.authenticate(senha)

            for i, pagina in enumerate(doc):
                if i >= 100:
                    logger.warning("PDF com mais de 100 páginas; OCR limitado a 100 (fitz).")
                    break
                try:
                    pix = pagina.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                    imagem = Image.open(io.BytesIO(pix.tobytes("png")))
                    imagens.append(imagem)
                except Exception as e_pag:
                    logger.warning("Erro ao renderizar página %d com PyMuPDF: %s", i + 1, e_pag)
        finally:
            doc.close()

        return imagens

    def _preprocessar_imagem(self, imagem: "Image.Image") -> "Image.Image":
        """
        Aplica pré-processamento na imagem para melhorar a qualidade do OCR.

        Passos:
          1. Converte para escala de cinza
          2. Aumenta o contraste
          3. Aplica nitidez
          4. Redimensiona para DPI ideal

        Args:
            imagem: Objeto PIL Image

        Returns:
            Imagem processada
        """
        # Escala de cinza
        imagem = imagem.convert("L")

        # Aumenta contraste
        enhancer = ImageEnhance.Contrast(imagem)
        imagem = enhancer.enhance(2.0)

        # Aumenta nitidez
        imagem = imagem.filter(ImageFilter.SHARPEN)

        return imagem

    def detectar_banco(self, texto: str) -> Optional[str]:
        """
        Tenta identificar o banco de origem pelo conteúdo do texto.

        Args:
            texto: Texto bruto do extrato

        Returns:
            Nome do banco identificado ou None
        """
        from app.utils.regex_patterns import BANCO_KEYWORDS

        texto_lower = texto.lower()
        for banco, keywords in BANCO_KEYWORDS.items():
            for kw in keywords:
                if kw in texto_lower:
                    return banco

        return None
